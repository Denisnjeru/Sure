import datetime
import os
from copy import copy
from pathlib import Path
import math
import requests
from django import apps
from django.shortcuts import get_object_or_404
from docx import Document
from openpyxl import Workbook, load_workbook
from sentry_sdk import capture_exception

# from apps.tender.models import QualityAssuranceResponse, QualityAssuranceQuestion
from apps.core.models import CategoryTypeSupplierLocation
from apps.core.utils import delete_matching_files_in_directory
from backend.storage_backends import PrivateMediaStorage


def get_file(file_url_path):
    A = PrivateMediaStorage()
    headers = {"ResponseContentDisposition": f"attachment;"}
    time = datetime.datetime.now()
    file_url = A.url(
        f"{file_url_path}", expire=300, parameters=headers, http_method="GET"
    )
    dir_name = Path(
        "media/temp/{}/{}/{}".format(time.year, time.month, time.day)
    )  # folder structure
    dir_name.mkdir(parents=True, exist_ok=True)
    file_name = os.path.basename(f"{file_url_path}")
    filepath = "{}/{}".format(dir_name, file_name)
    r = requests.get(file_url)
    with open("{}".format(filepath), "wb") as f:
        f.write(r.content)
    return filepath


def supplier_response_files(instance, filename):
    year = datetime.datetime.now().year
    tender = instance.question.section.category.tender
    return "%s/%s/%s/%s/%s/%s" % (
        tender.company.company_name.replace(" ", "_"),
        tender.unique_reference,
        instance.question.section.category.unique_reference,
        instance.supplier.company_name.replace(" ", "_"),
        year,
        filename,
    )


def supplier_letters(instance, filename):
    # file will be uploaded to MEDIA_ROOT/company_name/job_code/category_code/supplier_company_name/filename
    return "%s/%s/%s/%s/%s" % (
        instance.category.tender.company.company_name.replace(" ", "_"),
        instance.category.tender.unique_reference,
        instance.category.unique_reference,
        instance.supplier.company_name.replace(" ", "_"),
        filename,
    )


def c_suppliers(category, interim=False):
    supplier_category_score = apps.apps.get_model('tender', 'SupplierTechnicalScore')
    if interim:
        suppliers = (
            supplier_category_score.objects.filter(category_id=category.id)
            .order_by("rank")
            .only("supplier")
        )
    else:
        suppliers = (
            supplier_category_score.objects.filter(category_id=category.id)
            .order_by("rank_after_qa")
            .only("supplier")
        )
    return [s.supplier for s in suppliers]


def get_locations(category, supplier_id):
    supplier = get_object_or_404(apps.apps.get_model('suppliers', 'Supplier'), pk=supplier_id)

    sections = apps.apps.get_model('tender', 'Section').objects.filter(category_id=category.id)

    company_profile = sections.filter(name="Company Profile").first()

    if company_profile is not None:
        section_questions = company_profile.questions

        results = {
            "main_location": "", "location_1": "N/A", "location_2": "N/A",
            "location_3": "N/A", "location_4": "N/A", "location_5": "N/A",
            "location_6": "N/A",
        }

        for question in section_questions:
            if (
                question.short_description
                == "Physical address of business i.e. Town, street, building, Floor"
                or question.short_description
                == "5. Physical address of business - Main Street"
            ):
                results["main_location"] = supplier.tender_category_question_response(question)
            if question.short_description == "Main office location:":
                results["location_1"] = supplier.tender_category_question_response(question)
            if question.short_description == "Branch offices (if any):":
                results["location_2"] = supplier.tender_category_question_response(question)

        acf_ids = [240, 230, 231, 479]
        if category.tender_id in acf_ids:
            if category.tender_id == 230:
                """
                ACF Kenya
                """

                kenya_locations = sections.filter(name="Location").first()
                if kenya_locations is not None:
                    location_questions = kenya_locations.questions

                    for question in location_questions:
                        if (
                            question.short_description
                            == "Are you able to provide services in Mandera?"
                        ):
                            if supplier.tender_category_question_response(question) == "Yes":
                                results["location_2"] = "Mandera"
                            else:
                                results["location_2"] = "N/A"
                        if (
                            question.short_description
                            == "Are you able to provide services in Isiolo?"
                        ):
                            if supplier.tender_category_question_response(question) == "Yes":
                                results["location_3"] = "Isiolo"
                            else:
                                results["location_3"] = "N/A"
                        if (
                            question.short_description
                            == "Are you able to provide services in Westpokot?"
                        ):
                            if supplier.tender_category_question_response(question) == "Yes":
                                results["location_4"] = "Westpokot"
                            else:
                                results["location_4"] = "N/A"
                else:
                    results["location_1"] = "N/A"
                    results["location_2"] = "N/A"
                    results["location_3"] = "N/A"
                    results["location_4"] = "N/A"
                    results["location_5"] = "N/A"
                    results["location_6"] = "N/A"

            elif category.tender_id == 240:
                """
                ACF UG
                """
                uganda_locations = sections.filter(name="Location").first()
                if uganda_locations is not None:
                    location_questions = uganda_locations.questions

                for question in location_questions:
                    if (
                        question.short_description
                        == "1. Are you able to deliver directly to ACF-UG  offices Adjumani offices"
                    ):
                        if supplier.tender_category_question_response(question) == "Yes":
                            results["location_2"] = "Adjumani"
                        else:
                            results["location_2"] = "N/A"
                    if (
                        question.short_description
                        == "Are you able to deliver directly to ACF-UG  offices Yumbe offices"
                    ):
                        if supplier.tender_category_question_response(question) == "Yes":
                            results["location_3"] = "Yumbe"
                        else:
                            results["location_3"] = "N/A"
                    if (
                        question.short_description
                        == "Are you able to deliver directly to ACF-UG  offices Nakivale offices"
                    ):
                        if supplier.tender_category_question_response(question) == "Yes":
                            results["location_4"] = "Nakivale"
                        else:
                            results["location_4"] = "N/A"
                    if (
                        question.short_description
                        == "Are you able to deliver directly to ACF-UG  offices Kyangwali offices"
                    ):
                        if supplier.tender_category_question_response(question) == "Yes":
                            results["location_5"] = "Kyangwali"
                        else:
                            results["location_5"] = "N/A"
                    if (
                        question.short_description
                        == "Are you able to deliver directly to ACF-UG  offices Kiryandongo offices"
                    ):
                        if supplier.tender_category_question_response(question) == "Yes":
                            results["location_6"] = "Kiryandongo"
                        else:
                            results["location_6"] = "N/A"

            elif category.tender_id == 231:
                """
                ACF TZ
                """
                tz_locations = sections.filter(name="Locations").first()
                if tz_locations is not None:
                    location_questions = tz_locations.questions
                for question in location_questions:
                    if (
                        question.short_description
                        == "Are you able to supply countrywide?"
                    ):
                        if supplier.tender_category_question_response(question) == "Yes":
                            results["location_1"] = "Countrywide"
                        else:
                            results["location_1"] = "N/A"
                    if (
                        question.short_description
                        == "Are you able to supply in Dodoma city?"
                    ):
                        if supplier.tender_category_question_response(question) == "Yes":
                            results["location_2"] = "Dodoma"
                        else:
                            results["location_2"] = "N/A"
                    if (
                        question.short_description
                        == "Are you able to supply in Dar es salaam City?"
                    ):
                        if supplier.tender_category_question_response(question) == "Yes":
                            results["location_3"] = "Dar es Salaam"
                        else:
                            results["location_3"] = "N/A"
                    if (
                        question.short_description
                        == "Are you able to supply in Mpwapwa?"
                    ):
                        if supplier.tender_category_question_response(question) == "Yes":
                            results["location_4"] = "Mpwapwa"
                        else:
                            results["location_4"] = "N/A"
                    if (
                        question.short_description
                        == "Are you able to supply in Itigi District Council?"
                    ):
                        if supplier.tender_category_question_response(question) == "Yes":
                            results["location_4"] = "Itigi"
                        else:
                            results["location_4"] = "N/A"
                    if (
                        question.short_description
                        == "Are you able to supply in Singida Region?"
                    ):
                        if supplier.tender_category_question_response(question) == "Yes":
                            results["location_5"] = "Singida"
                        else:
                            results["location_5"] = "N/A"
                    if (
                        question.short_description
                        == "Are you able to supply in Iramba District?"
                    ):
                        if supplier.tender_category_question_response(question) == "Yes":
                            results["location_6"] = "Iramba"
                        else:
                            results["location_6"] = "N/A"

            elif category.tender_id == 479:
                """
                ACF Ethiopia
                """

                et_locations = sections.filter(name="Location").first()
                if et_locations is not None:
                    location_questions = et_locations.questions
                for question in location_questions:
                    if (
                        question.short_description
                        == "1. Are you able to provide services or supply in Addis Ababa?"
                    ):
                        if supplier.tender_category_question_response(question) == "Yes":
                            results["location_1"] = "Addis Ababa"
                        else:
                            results["location_1"] = "N/A"
                    if (
                        question.short_description
                        == "2. Are you able to provide services or supply in Gambella?"
                    ):
                        if supplier.tender_category_question_response(question) == "Yes":
                            results["location_2"] = "Gambela"
                        else:
                            results["location_2"] = "N/A"
                    if (
                        question.short_description
                        == "3. Are you able to provide services or supply in Yabello?"
                    ):
                        if supplier.tender_category_question_response(question) == "Yes":
                            results["location_3"] = "Yabello"
                        else:
                            results["location_3"] = "N/A"
                    if (
                        question.short_description
                        == "4. Are you able to provide services or supply in Bahir Dar?"
                    ):
                        if supplier.tender_category_question_response(question) == "Yes":
                            results["location_4"] = "Bahir Dar"
                        else:
                            results["location_4"] = "N/A"
                    if (
                        question.short_description
                        == "5. Are you able to provide services or supply in Harar?"
                    ):
                        if supplier.tender_category_question_response(question) == "Yes":
                            results["location_4"] = "Harar"
                        else:
                            results["location_4"] = "N/A"
                    if (
                        question.short_description
                        == "6. Are you able to provide services or supply in Gode?"
                    ):
                        if supplier.tender_category_question_response(question) == "Yes":
                            results["location_5"] = "Gode"
                        else:
                            results["location_5"] = "N/A"
                    if (
                        question.short_description
                        == "7. Are you able to provide services or supply in Mekelle?"
                    ):
                        if supplier.tender_category_question_response(question) == "Yes":
                            results["location_6"] = "Mekelle"
                        else:
                            results["location_6"] = "N/A"
                    if (
                        question.short_description
                        == "8. Are you able to provide services or supply in Gimbi?"
                    ):
                        if supplier.tender_category_question_response(question) == "Yes":
                            results["location_7"] = "Gimbi"
                        else:
                            results["location_7"] = "N/A"

        # jobs/companies with location question in SLA
        sla_ids = [499, 500]
        if category.tender_id in sla_ids:
            if category.tender_id == 499 or category.tender_id == 500:
                """
                CITAM
                """

                locations = sections.filter(
                    name="Service Level Requirement"
                ).first()
                if locations is not None:
                    location_questions = locations.questions

                    for question in location_questions:
                        if (
                            question.description
                            == "2. Do you have presence in Kisumu?"
                        ):
                            if supplier.tender_category_question_response(question) == "Yes":
                                results["location_1"] = "Kisumu"
                            else:
                                results["location_1"] = "N/A"
                        if (
                            question.description
                            == "3. Do you have presence in Eldoret?"
                        ):
                            if supplier.tender_category_question_response(question) == "Yes":
                                results["location_2"] = "Eldoret"
                            else:
                                results["location_2"] = "N/A"
                        if (
                            question.description
                            == "4. Do you have presence in Kapsabet?"
                        ):
                            if supplier.tender_category_question_response(question) == "Yes":
                                results["location_3"] = "Kapsabet"
                            else:
                                results["location_3"] = "N/A"

                        if (
                            question.description
                            == "5. Do you have presence in Nakuru?"
                        ):
                            if supplier.tender_category_question_response(question) == "Yes":
                                results["location_4"] = "Nakuru"
                            else:
                                results["location_4"] = "N/A"

                        if (
                            question.description
                            == "6. Do you have presence in Naivasha?"
                        ):
                            if supplier.tender_category_question_response(question) == "Yes":
                                results["location_5"] = "Naivasha"
                            else:
                                results["location_5"] = "N/A"
                        if (
                            question.description
                            == "7. Do you have presence in Nairobi?"
                        ):
                            if supplier.tender_category_question_response(question) == "Yes":
                                results["location_6"] = "Nairobi"
                            else:
                                results["location_6"] = "N/A"
                        if (
                            question.description
                            == "8. Do you have presence in Kisii?"
                        ):
                            if supplier.tender_category_question_response(question) == "Yes":
                                results["location_7"] = "Kisii"
                            else:
                                results["location_7"] = "N/A"
                        if (
                            question.description
                            == "9. Do you have presence in Meru?"
                        ):
                            if supplier.tender_category_question_response(question) == "Yes":
                                results["location_8"] = "Meru"
                            else:
                                results["location_8"] = "N/A"
                        if (
                            question.description
                            == "10. Do you have presence in Nyeri?"
                        ):
                            if supplier.tender_category_question_response(question) == "Yes":
                                results["location_9"] = "Nyeri"
                            else:
                                results["location_9"] = "N/A"
                        if (
                            question.description
                            == "11. Do you have presence in Mombasa?"
                        ):
                            if supplier.tender_category_question_response(question) == "Yes":
                                results["location_10"] = "Mombasa"
                            else:
                                results["location_10"] = "N/A"
                else:
                    results["location_1"] = "N/A"
                    results["location_2"] = "N/A"
                    results["location_3"] = "N/A"
                    results["location_4"] = "N/A"
                    results["location_5"] = "N/A"
                    results["location_6"] = "N/A"
                    results["location_7"] = "N/A"
                    results["location_8"] = "N/A"
                    results["location_9"] = "N/A"
                    results["location_10"] = "N/A"
        return results

    context = {"status": 400, "message": "Tender does not exist"}
    return context


def create_award_letter(letter_id):
    letter = get_object_or_404(apps.apps.get_model('tender', 'AwardLetter'), pk=letter_id)

    time = datetime.datetime.now()
    dir_name = Path(
        "media/award_letters/{}/{}".format(time.year, time.month)
    )  # folder structure
    dir_name.mkdir(parents=True, exist_ok=True)
    time_only = "%d%d%d%d" % (time.day, time.hour, time.minute, time.second)
    filepath = "{}/{}_{}_{}.docx".format(
        dir_name,
        letter.category.prequalification.id,
        letter.supplier.company_name.replace(" ", "_") + "Award_Letter",
        time_only,
    )

    document = Document()
    if letter.supplier.logo_url:
        document.add_picture(letter.supplier.logo_url)

    document.add_heading(letter.supplier.company_name, level=4)
    document.add_heading("%s" % letter.award_date.strftime("%B %d, %Y"), level=4)
    document.add_heading("Dear Sir/Madam,", level=4)
    document.add_heading(letter.category.name, level=4)

    paragraph_one = "The above matter refers."
    paragraph_two = (
        "Thank you for participating in the prequalification of suppliers for "
        + letter.category.name
        + "."
    )
    paragraph_three = (
        "After an evaluation of your bid under the above-mentioned category, we are pleased to inform "
        "you that your bid was successful. Your company has therefore been prequalified among our list "
        "of vendors under the stated category for the period {{period}}. This will be subject to our terms "
        "and conditions."
    )
    paragraph_four = (
        "This letter serves as a notification of the success of your application and should not be used as a "
        "basis of incurring any expense or entering into any binding arrangements with any third party "
        "on our behalf. We will request for quotations as and when necessary based on our needs."
    )
    paragraph_five = (
        "Please confirm receipt and acceptance of this communication by signing and stamping this letter "
        f"and sending to us via email to {letter.supplier.company_name}."
    )
    paragraph_six = "Thank you."
    paragraph_seven = "Yours faithfully,"
    paragraph_eight = "............................."
    paragraph_nine = "............................."
    document.add_paragraph(paragraph_one)
    document.add_paragraph(paragraph_two)
    document.add_paragraph(paragraph_three)
    document.add_paragraph(paragraph_four)
    document.add_paragraph(paragraph_five)
    document.add_paragraph(paragraph_six)
    document.add_paragraph(paragraph_seven)
    document.add_paragraph(paragraph_eight)
    document.add_paragraph(paragraph_nine)
    document.add_page_break()
    document.save(filepath)

    time = datetime.datetime.now()
    try:
        with open(filepath, "rb") as l:
            storage = PrivateMediaStorage()
            url = f"award_letters/{time.year}/{time.month}/{time.day}/{os.path.basename(l.name)}"
            storage.save(url, l)

            letter.letter = url
            letter.description = document
            letter.save()
            context = {
                "filepath": url,
                "description": document
            }
            return context
    except Exception as e:
        print(e)
        return


def create_regret_letter(letter_id):
    letter = get_object_or_404(apps.apps.get_model('tender', 'RegretLetter'), pk=letter_id)

    time = datetime.datetime.now()
    dir_name = Path(
        "media/regret_letters/{}/{}".format(time.year, time.month)
    )  # folder structure
    dir_name.mkdir(parents=True, exist_ok=True)
    time_only = "%d%d%d%d" % (time.day, time.hour, time.minute, time.second)
    filepath = "{}/{}_{}_{}.docx".format(
        dir_name,
        letter.category.prequalification.id,
        letter.supplier.company_name.replace(" ", "_") + "Regret_Letter",
        time_only,
    )

    document = Document()
    if letter.supplier.logo_url:
        document.add_picture(letter.supplier.logo_url)

    document.add_heading(letter.supplier.company_name, level=4)
    document.add_heading("%s" % letter.regret_date.strftime("%B %d, %Y"), level=4)
    document.add_heading("Dear Sir/Madam,", level=4)
    document.add_heading(letter.category.name, level=4)

    paragraph_one = "The above matter refers."
    paragraph_two = (
        "Thank you for participating in the prequalification of suppliers for "
        + letter.category.name
        + "."
    )
    paragraph_three = (
        "Following a detailed evaluation of your bid, we regret to inform you that on this occasion "
        "your application has not been successful."
    )
    paragraph_four = (
        "We however thank you for your interest to work with "
        + letter.category.prequalification.company.company_name
        + " and look forward to "
        "working with you in the near future. Your records have been maintained and may call upon "
        "you if need arises."
    )
    paragraph_five = "Thank you."
    paragraph_six = "Yours faithfully,"
    paragraph_seven = "............................."
    paragraph_eight = "............................."
    document.add_paragraph(paragraph_one)
    document.add_paragraph(paragraph_two)
    document.add_paragraph(paragraph_three)
    document.add_paragraph(paragraph_four)
    document.add_paragraph(paragraph_five)
    document.add_paragraph(paragraph_six)
    document.add_paragraph(paragraph_seven)
    document.add_paragraph(paragraph_eight)

    document.add_page_break()
    document.save(filepath)

    time = datetime.datetime.now()
    try:
        with open(filepath, "rb") as l:
            storage = PrivateMediaStorage()
            url = f"regret_letters/{time.year}/{time.month}/{time.day}/{os.path.basename(l.name)}"
            storage.save(url, l)

            letter.letter = url
            letter.description = document
            letter.save()
            context = {
                "filepath": url,
                "description": document
            }
            return context
    except Exception as e:
        print(e)
        return


def get_supplier_data(category_id, supplier_id):
    try:
        category = apps.apps.get_model('tender', 'Category').objects.filter(id=category_id).first()
        supplier = apps.apps.get_model('suppliers', 'Supplier').objects.filter(id=supplier_id).first()

        section_data = {
            "Bank_Details": "",
            "Office_Address": "",
            "P.O Box address": "",
            "Date_Incorporation": "",
            "Main_Office": "",
            "Branch_Office": "",
        }
        registration_section = category.sections.filter(
            name="Registration and Statutory  Requirements"
        ).first()
        company_profile = category.sections.filter(name="Company Profile").first()
        # registration section questions
        registration_qn = registration_section.questions.first()

        date_incorporation = get_date_of_incorporation(supplier, registration_qn)
        section_data["Date_Incorporation"] = date_incorporation

        for question in registration_section.questions:
            if question.short_description == "Company registration certificate number":
                section_data["Co_Reg_No"] = supplier.tender_category_question_response(question)
            if question.short_description == "Bank Details":
                section_data["Bank_Details"] = supplier.tender_category_question_response(question)

        # company profile section questions
        for question in company_profile.questions:
            if (
                question.short_description
                == "Physical address of business i.e. Town, street, building, Floor"
            ):
                section_data["Office_Address"] = supplier.tender_category_question_response(question)
            if question.short_description == "P.O Box address":
                section_data["P.O Box address"] = supplier.tender_category_question_response(question)
            if question.short_description == "Main office location:":
                section_data["Main_Office"] = supplier.tender_category_question_response(question)
            if question.short_description == "Branch offices (if any):":
                section_data["Branch_Office"] = supplier.tender_category_question_response(question)

        supplier_data = {
            "Company_Name": supplier.company_name,
            "Contact_Person": supplier.contact_name,
            "Phone_Number": supplier.phone_number,
            "Email": supplier.email,
            "Tax_ID": supplier.kra_pin_number,
            "Category_Code": category.unique_reference,
            "Category": category.name,
            "Postal_Address": supplier.address,
            "Location": supplier.location,
            "Section": section_data,
        }

        return supplier_data

    except Exception as e:
        capture_exception(e)
        print(e)


def get_date_of_incorporation(supplier, question):
    """
    Get the date of incorporation from qa reponses
    """
    try:
        qa_res = apps.apps.get_model('tender', 'QualityAssuranceResponse').objects.filter(
            supplier=supplier,
            quality_assurance_question=apps.apps.get_model('tender', 'QualityAssuranceQuestion').objects.filter(
                question=question
            ).first(),
        ).first()
        if qa_res is not None:
            date_incorporation = qa_res.date
            date_incorporation.date().strftime("%Y-%m-%d")
        else:
            date_incorporation = ""

        return date_incorporation

    except Exception as e:
        capture_exception(e)


# Category Type Suppliers
def get_cat_supplier_locations(cat_type_supplier):
    supplier_locations = CategoryTypeSupplierLocation.objects.filter(
        category_type_supplier=cat_type_supplier
    )
    locations = []
    if len(supplier_locations) > 1:
        for i, loc in enumerate(supplier_locations, start=1):
            locs = {f"location_{i}": loc.location}
            locations.append(locs)

    locations_headers = []
    try:
        if len(locations) > 1:
            for i, loc in enumerate(locations, start=1):
                locations_headers.append(loc[f"location_{i}"])
        else:
            locations_headers.append("N/A")

        return locations_headers

    except Exception as e:
        capture_exception(e)


def tender_criteria_template_download(job):
    time = datetime.datetime.now()
    dir_name = Path(
        "media/criteria_templates/{}/{}".format(time.year, time.month)
    )  # folder structure
    dir_name.mkdir(parents=True, exist_ok=True)

    # time_only = "%d%d%d%d" % (time.day, time.hour, time.minute, time.second)
    filepath = "{}/{}_{}_Questions_Template.xlsx".format(
        dir_name, job.id, job.company.company_name.replace(" ", "_")
    )

    match_string = "{}_{}_Questions_Template".format(
        job.id, job.company.company_name.replace(" ", "_")
    )
    delete_matching_files_in_directory(dir_name, match_string, ".xlsx")

    categories = apps.apps.get_model('tender', 'Category').objects.filter(tender_id=job.id)
    workbook = Workbook()
    columns_to_hide = ['A', 'C', 'D', 'F', 'G', 'H', 'I']

    for category in categories:
        category_type = category.category_type
        criteria = apps.apps.get_model('core', 'CategoryTypeCriteria').objects.filter(
            category_type_id=category_type.id, criteria_country_id=job.criteria_country_id
        ).first()
        if criteria:
            workbook.create_sheet(f"{category.unique_reference}")
            criteria_sheet = workbook[f"{category.unique_reference}"]
            try:
                category_type_workbook = load_workbook(criteria.file_url)
                category_type_worksheet = category_type_workbook[f'{category_type.innitials}']

                for row in category_type_worksheet:
                    for cell in row:
                        quote_cell = criteria_sheet[cell.coordinate]
                        quote_cell.value = cell.value
                        if cell.has_style:
                            quote_cell.font = copy(cell.font)
                            quote_cell.border = copy(cell.border)
                            quote_cell.fill = copy(cell.fill)
                            quote_cell.number_format = copy(cell.number_format)
                            quote_cell.protection = copy(cell.protection)
                            quote_cell.alignment = copy(cell.alignment)

                for col in columns_to_hide:
                    criteria_sheet.column_dimensions[col].hidden = True
            except Exception as e:
                capture_exception(e)
                continue
        else:
            workbook.create_sheet(f"{job.unique_reference}_{category.unique_reference}")
    workbook.remove(workbook['Sheet'])
    workbook.save(filepath)
    return filepath

def send_current_suppliers_letter_email(job_id):
    """
    Sends job advert and letter to current suppliers 
    :param tender_id:
    :return:
    """
    job = apps.apps.get_model('tender', 'Tender').objects.filter(id=job_id).first()
    open_categories = list(apps.apps.get_model('tender', 'Category').objects.filter(prequalification_id=job_id, is_open=True).values_list("id", flat=True))
    company = job.company
    current_suppliers = (
        apps.apps.get_model('core', 'CurrentSupplier').objects.filter(company=company, category_id__in=open_categories, target=ContentType.objects.get_for_model(apps.apps.get_model('tender', 'Category')))
        .order_by("supplier_email").distinct("supplier_email")
    )
    category = apps.apps.get_model('tender', 'Category').objects.filter(prequalification_id=job_id).first()

    supplier_emails_list = list(current_suppliers.values_list("supplier_email", flat=True))
    # split the email list into chunks of 20 emails
    n = 20
    supplier_emails_group = [supplier_emails_list[i:i + n] for i in range(0, len(supplier_emails_list), n)]
 
    email_subject = _(
                "Tendersure: %s "
                % job.title
            )
    body = render_to_string(
        "tender/letters/current_suppliers_email.html",
        {"company_name": company.company_name, "closing_date": category.closing_date},
    )
    message = render_to_string(
        "tender/letters/current_suppliers_email.html",
        {"company_name": company.company_name, "closing_date": category.closing_date},
    )
    for group in supplier_emails_group:
        email = EmailMultiAlternatives(
            subject=email_subject, body=body, to=["tendersure@mail.com"], bcc=group
        )
        email.attach_alternative(message, "text/html")       
        try:
            email.attach_file(get_file(job.advert))
        except:
            pass
        try:
            email.attach_file(get_file(job.current_suppliers_letter))
        except:
            pass
        email.send()
    

def calculate_financial_ratios_after_qa(instance_id, user):
    """
    Calculate Ratios After QA
    debt/equity = long term loans/equity
    current ratio = current assets/current liabilities (current assets+ cash)/current liabilities
    cash ratio = cash/current liabilities
    GP margin = GP/turnover
    NP margin = NP/Turnover
    """
    ratio_instance = apps.apps.get_model('tender', 'FinancialRatio').objects.filter(id=instance_id).first()
    section = ratio_instance.section

    all_questions = apps.apps.get_model('tender', 'Question').objects.filter(
        section_id=ratio_instance.section_id).order_by("id")

    first_question = all_questions.first()
    other_questions = all_questions.exclude(id=first_question.id)
    supplier = ratio_instance.supplier

    section_score = 0
    try:
        debt_equity_ratio = float(ratio_instance.debtors_after_qa) / float(ratio_instance.equity_after_qa)
    except:
        debt_equity_ratio = 0
    try:
        current_ratio = float(ratio_instance.current_assets_after_qa) / float(ratio_instance.curr_liabilities_after_qa)
    except:
        current_ratio = 0

    try:
        cash_ratio = float(ratio_instance.cash_after_qa) / float(ratio_instance.curr_liabilities_after_qa)
    except:
        cash_ratio = 0

    try:
        gp_margin = (
            float(ratio_instance.gross_profit_after_qa) / float(ratio_instance.turnover_after_qa)) * 100
    except:
        gp_margin = 0

    try:
        np_margin = (
            float(ratio_instance.net_profit_after_qa)/ float(ratio_instance.turnover_after_qa)) * 100
    except:
        np_margin = 0

    ratio_list = [debt_equity_ratio, current_ratio, cash_ratio, gp_margin, np_margin]

    ac_question = apps.apps.get_model('tender', 'Question').objects.filter(
        description="1. For limited liability companies, attach audited accounts for the last two years, "
                    "for sole proprietors and partnerships, attach your most recent management accounts",
        section__category_id=section.category_id
    ).first()

    if ac_question is not None:
        supplier_response = apps.apps.get_model('tender', 'SupplierResponse').objects.filter(
            question_id=ac_question.id, supplier_id=supplier.id).first()
    else:
        supplier_response = None

    # create QAResponse
    c = 0
    for question in other_questions:
        ratio = ratio_list[c]
        question_options = question.options
        c += 1
        data = None
        if ratio is not None:
            for option in question_options:
                start = float(option.split("-", 1)[0])
                end = float(option.split("-", 1)[1])
                q_ratio = math.floor(float(ratio) * 10)/10

                if start <= q_ratio <= end:
                    my_index = question_options.index(option)
                    score_after_qa = float(question.scores[my_index])
                    data = {
                        "supplier": supplier, "question": question,
                        "score_after_qa": score_after_qa, "outcome": "Pass"
                    }
                    break
                else:
                    continue

        if data is None:
            data = {
                "supplier": supplier, "question": question,
                "score_after_qa": 0, "outcome": "Pass"
            }

        """ Check if supplier submitted their financials file if not give zero marks """
        if supplier_response is None:
            data = {
                "supplier": supplier, "question": question,
                "score_after_qa": 0, "outcome": "Fail"
            }
        elif supplier_response is not None:
            if supplier_response.options is None and supplier_response.document_url is None or\
                    len(supplier_response.options) < 1 and not supplier_response.document_url.name:
                data = {
                    "supplier": supplier, "question": question,
                    "score_after_qa": 0, "outcome": "Fail"
                }

        qa_question = apps.apps.get_model('tender', 'QualityAssuranceQuestion').objects.filter(question=question).first()

        if qa_question is not None:
            qa_res = apps.apps.get_model('tender', 'QualityAssuranceResponse').objects.filter(
                supplier=data["supplier"], quality_assurance_question=qa_question).last()

            if qa_res is not None:
                qa_res.score_after_qa = data["score_after_qa"]
                qa_res.created_by = user
                qa_res.number = ratio
                qa_res.outcome = data["outcome"]
                qa_res.save()
            else:
                qa_res = apps.apps.get_model('tender', 'QualityAssuranceResponse').objects.create(
                    supplier=data["supplier"], quality_assurance_question=qa_question,
                    score_after_qa=data["score_after_qa"], created_by=user, outcome=data["outcome"]
                )
            context = {"qa_res": qa_res}
        else:
            context = {}

        print(f"Context: {context}")
    # prequal = apps.apps.get_model('prequal', 'Prequalification').objects.filter(category_id=section.category_id).first()
    # supplier.resolve_qa_scores(prequal)
    return context


def ratio_scores_before_after_qa(supplier_id, section_id):
    section = apps.apps.get_model("tender", "Section").objects.filter(id=section_id).first()
    supplier = apps.apps.get_model("suppliers", "Supplier").objects.filter(id=supplier_id).first()

    if section is not None and supplier is not None:
        questions = section.questions.order_by("id")[1:]
        questions_data = []
        c = 0
        for question in questions:
            before_qa = supplier.tender_question_score(question)
            after_qa = 0
            question_in_qa = apps.apps.get_model("tender", "QualityAssuranceResponse").objects.filter(
                quality_assurance_question__question=question, supplier__id=supplier.id).last()
            if question_in_qa:
                if question_in_qa.score_after_qa is not None:
                    after_qa = question_in_qa.score_after_qa

            data = {
                "index": c,
                "before_qa": before_qa,
                "after_qa": float(after_qa),
                "variance": float(after_qa) - float(before_qa),
            }
            questions_data.append(data)
            c += 1
    
    return questions_data
            
                
def variance_ratios(instance_id, *args, **kwargs):
    ratio_instance = apps.apps.get_model("tender", "FinancialRatio").objects.filter(id=instance_id).first()
    if ratio_instance is not None:
        equity = ratio_instance.equity if ratio_instance.equity is not None else 0
        equity_qa = ratio_instance.equity_after_qa if ratio_instance.equity_after_qa is not None else 0
        if equity >= 0:
            eq_variance = equity_qa - equity
        else:
            eq_variance = equity + equity_qa
        eq_comment = "Pass" if eq_variance == 0 else "Fail"

        debtors = ratio_instance.debtors if ratio_instance.debtors is not None else 0
        debtors_qa = ratio_instance.debtors_after_qa if ratio_instance.debtors_after_qa is not None else 0
        d_variance = debtors_qa - debtors
        d_comment = "Pass" if d_variance == 0 else "Fail"

        curr_liabilities = ratio_instance.curr_liabilities if ratio_instance.curr_liabilities is not None else 0
        curr_liabilities_qa = ratio_instance.curr_liabilities_after_qa if ratio_instance.curr_liabilities_after_qa is not None else 0
        curr_variance = curr_liabilities_qa - curr_liabilities
        curr_comment = "Pass" if curr_variance == 0 else "Fail"

        f_assets = ratio_instance.fixed_assets if ratio_instance.fixed_assets is not None else 0
        f_assets_qa = ratio_instance.fixed_assets_after_qa if ratio_instance.fixed_assets_after_qa is not None else 0
        f_variance = f_assets_qa - f_assets
        f_comment = "Pass" if f_variance == 0 else "Fail"

        c_assets = ratio_instance.current_assets if ratio_instance.current_assets is not None else 0
        c_assets_qa = ratio_instance.current_assets_after_qa if ratio_instance.current_assets_after_qa is not None else 0
        c_variance = c_assets_qa - c_assets
        c_comment = "Pass" if c_variance == 0 else "Fail"

        cash = ratio_instance.cash if ratio_instance.cash is not None else 0
        cash_qa = ratio_instance.cash_after_qa if ratio_instance.cash_after_qa is not None else 0
        cash_variance = cash_qa - cash
        cash_comment = "Pass" if cash_variance == 0 else "Fail"

        turnover = ratio_instance.turnover if ratio_instance.turnover is not None else 0
        turnover_qa = ratio_instance.turnover_after_qa if ratio_instance.turnover_after_qa is not None else 0
        turnover_variance = turnover_qa - turnover
        turnover_comment = "Pass" if turnover_variance == 0 else "Fail"

        g_profit = ratio_instance.gross_profit if ratio_instance.gross_profit is not None else 0
        g_profit_qa = ratio_instance.gross_profit_after_qa if ratio_instance.gross_profit_after_qa is not None else 0
        g_variance = g_profit_qa - g_profit
        g_comment = "Pass" if g_variance == 0 else "Fail"

        n_profit = ratio_instance.net_profit if ratio_instance.net_profit is not None else 0
        n_profit_qa = ratio_instance.net_profit_after_qa if ratio_instance.net_profit_after_qa is not None else 0
        n_variance = n_profit_qa - n_profit
        n_comment = "Pass" if n_variance == 0 else "Fail"

        data = {
            "equity":{"variance": eq_variance,"comment":eq_comment},
            "debtors":{"variance":d_variance, "comment":d_comment},
            "curr_liabilities": {"variance":curr_variance,"comment":curr_comment},
            "f_assets": {"variance":f_variance, "comment":f_comment},
            "c_assets": {"variance":c_variance, "comment":c_comment},
            "cash": {"variance":cash_variance, "comment":cash_comment},
            "turnover": {"variance":turnover_variance, "comment":turnover_comment},
            "g_profit": {"variance":g_variance, "comment":g_comment},
            "n_profit": {"variance":n_variance, "comment":n_comment},
        } 

        return data


def ratios_before_after_qa(instance_id):
    ratio_instance = apps.apps.get_model("tender", "FinancialRatio").objects.filter(id=instance_id).first()
    section = ratio_instance.section

    try:
        debt_equity_ratio_qa = float(ratio_instance.debtors_after_qa) / float(
            ratio_instance.equity_after_qa
        )
    except:
        debt_equity_ratio_qa = 0
    try:
        current_ratio_qa = float(ratio_instance.current_assets_after_qa) / float(ratio_instance.curr_liabilities_after_qa)
    except:
        current_ratio_qa = 0

    try:
        cash_ratio_qa = float(ratio_instance.cash_after_qa) / float(
            ratio_instance.curr_liabilities_after_qa
        )
    except:
        cash_ratio_qa = 0

    try:
        gp_margin_qa = (
            float(ratio_instance.gross_profit_after_qa) / float(ratio_instance.turnover_after_qa)) * 100
    except:
        gp_margin_qa = 0

    try:
        np_margin_qa = (
            float(ratio_instance.net_profit_after_qa) / float(ratio_instance.turnover_after_qa)) * 100
    except:
        np_margin_qa = 0

    try:
        debt_equity_ratio = float(ratio_instance.debtors) / float(ratio_instance.equity)
    except:
        debt_equity_ratio = 0

    try:
        current_ratio = (float(ratio_instance.current_assets)) / float(ratio_instance.curr_liabilities)
    except:
        current_ratio = 0

    try:
        cash_ratio = float(ratio_instance.cash) / float(ratio_instance.curr_liabilities)
    except:
        cash_ratio = 0

    try:
        gp_margin = (
            float(ratio_instance.gross_profit) / float(ratio_instance.turnover)
        ) * 100
    except:
        gp_margin = 0

    try:
        np_margin = (
            float(ratio_instance.net_profit) / float(ratio_instance.turnover)
        ) * 100
    except:
        np_margin = 0

    data = {
        "debt_equity": {
            "before_qa": math.floor(float(debt_equity_ratio) * 100)/100,
            "after_qa": math.floor(float(debt_equity_ratio_qa) * 100)/100,
            "variance": math.floor(float(debt_equity_ratio_qa - debt_equity_ratio) * 100)/100,
        },
        "current_ratio": {
            "before_qa": math.floor(float(current_ratio) * 100)/100,
            "after_qa": math.floor(float(current_ratio_qa) * 100)/100,
            "variance": math.floor(float(current_ratio_qa - current_ratio) * 100)/100,
        },
        "cash_ratio": {
            "before_qa": math.floor(float(cash_ratio) * 100)/100,
            "after_qa": math.floor(float(cash_ratio_qa) * 100)/100,
            "variance": math.floor(float(cash_ratio_qa - cash_ratio) * 100)/100,
        },
        "gp_margin": {
            "before_qa": math.floor(float(gp_margin) * 100)/100,
            "after_qa": math.floor(float(gp_margin_qa) * 100)/100,
            "variance": math.floor(float(gp_margin_qa - gp_margin) * 100)/100,
        },
        "np_margin": {
            "before_qa": math.floor(float(np_margin) * 100)/100,
            "after_qa": math.floor(float(np_margin_qa) * 100)/100,
            "variance": math.floor(float(np_margin_qa - np_margin) * 100)/100,
        },
    }

    return data
